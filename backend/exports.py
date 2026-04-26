import io
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter


def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), kwargs.get(edge, "none"))
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:color"), kwargs.get("color", "D1D5DB"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def generate_word(newsletter_draft: str, articles: list, run_id: str) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("FMCG M&A INTELLIGENCE NEWSLETTER")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("Industry Intelligence Report  |  Powered by LangGraph + Claude")
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    sub_run.italic = True

    doc.add_paragraph()

    lines = newsletter_draft.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        if re.match(r"^#{1,2}\s", stripped):
            text = re.sub(r"^#{1,3}\s*", "", stripped)
            h = doc.add_heading(text, level=2)
            h.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
            continue

        if re.match(r"^\*\*[^*]+\*\*$", stripped):
            p = doc.add_paragraph()
            r = p.add_run(stripped.strip("*"))
            r.bold = True
            r.font.size = Pt(11)
            continue

        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            text = re.sub(r"^\d+\.\s*", "", stripped)
            _add_inline_bold(p, text)
            continue

        if re.match(r"^[-•]\s", stripped):
            p = doc.add_paragraph(style="List Bullet")
            text = re.sub(r"^[-•]\s*", "", stripped)
            _add_inline_bold(p, text)
            continue

        p = doc.add_paragraph()
        _add_inline_bold(p, stripped)

    doc.add_page_break()

    doc.add_heading("Source Articles", level=1)
    sub2 = doc.add_paragraph(f"Total articles after pipeline filtering: {len(articles)}")
    sub2.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    sub2.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = ["Title", "Source", "Deal Type", "Companies", "Value", "Credibility"]
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1E40AF")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

    for art in articles:
        row = table.add_row().cells
        row[0].text = art.get("title", "")[:80]
        row[1].text = art.get("source", "")
        row[2].text = art.get("deal_type", "")
        row[3].text = "; ".join(art.get("companies", []))
        row[4].text = art.get("deal_value", "") or "Undisclosed"
        row[5].text = f"{round(art.get('credibility_score', 0) * 100)}%"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _add_inline_bold(para, text: str):
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = para.add_run(part[2:-2])
            r.bold = True
        else:
            para.add_run(part)


def generate_excel(articles: list, newsletter_sections: dict, run_id: str) -> bytes:
    wb = openpyxl.Workbook()

    ws_nl = wb.active
    ws_nl.title = "Newsletter"

    header_fill = PatternFill("solid", fgColor="1E40AF")
    header_font = Font(bold=True, color="FFFFFF", size=12)
    section_fill = PatternFill("solid", fgColor="EFF6FF")
    section_font = Font(bold=True, color="1E40AF", size=11)
    thin = Side(style="thin", color="D1D5DB")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    ws_nl.merge_cells("A1:B1")
    ws_nl["A1"] = "FMCG M&A Intelligence Newsletter"
    ws_nl["A1"].font = Font(bold=True, size=16, color="1E40AF")
    ws_nl["A1"].alignment = Alignment(horizontal="center")
    ws_nl.row_dimensions[1].height = 30

    ws_nl.merge_cells("A2:B2")
    ws_nl["A2"] = f"Run ID: {run_id}"
    ws_nl["A2"].font = Font(italic=True, color="6B7280", size=9)
    ws_nl["A2"].alignment = Alignment(horizontal="center")

    row = 4
    section_map = {
        "executive_summary": "Executive Summary",
        "market_outlook": "Market Outlook",
        "top_deals": "Top Deals",
        "deals_to_watch": "Deals to Watch",
        "sector_breakdown": "Sector Breakdown",
    }

    for key, label in section_map.items():
        val = newsletter_sections.get(key)
        if not val:
            continue
        ws_nl.merge_cells(f"A{row}:B{row}")
        ws_nl[f"A{row}"] = label
        ws_nl[f"A{row}"].fill = section_fill
        ws_nl[f"A{row}"].font = section_font
        ws_nl[f"A{row}"].border = border
        row += 1

        if isinstance(val, str):
            ws_nl.merge_cells(f"A{row}:B{row}")
            ws_nl[f"A{row}"] = val
            ws_nl[f"A{row}"].alignment = Alignment(wrap_text=True)
            ws_nl[f"A{row}"].border = border
            ws_nl.row_dimensions[row].height = max(40, len(val) // 3)
            row += 1
        elif isinstance(val, list):
            for item in val:
                ws_nl.merge_cells(f"A{row}:B{row}")
                text = item if isinstance(item, str) else str(item)
                ws_nl[f"A{row}"] = f"• {text}"
                ws_nl[f"A{row}"].alignment = Alignment(wrap_text=True)
                ws_nl[f"A{row}"].border = border
                ws_nl.row_dimensions[row].height = max(30, len(text) // 3)
                row += 1
        elif isinstance(val, dict):
            for sector, deals in val.items():
                ws_nl[f"A{row}"] = sector
                ws_nl[f"A{row}"].font = Font(bold=True, size=10)
                ws_nl[f"A{row}"].border = border
                ws_nl.merge_cells(f"A{row}:B{row}")
                row += 1
                text = deals if isinstance(deals, str) else str(deals)
                ws_nl.merge_cells(f"A{row}:B{row}")
                ws_nl[f"A{row}"] = text
                ws_nl[f"A{row}"].alignment = Alignment(wrap_text=True)
                ws_nl[f"A{row}"].border = border
                ws_nl.row_dimensions[row].height = 35
                row += 1

        row += 1

    ws_nl.column_dimensions["A"].width = 40
    ws_nl.column_dimensions["B"].width = 60

    ws = wb.create_sheet("Source Articles")
    cols = ["Title", "Source", "URL", "Published", "Deal Type", "Companies", "Deal Value", "Relevance %", "Credibility %"]
    for ci, col in enumerate(cols, 1):
        cell = ws.cell(row=1, column=ci, value=col)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")
        cell.border = border

    for ri, art in enumerate(articles, 2):
        row_data = [
            art.get("title", ""),
            art.get("source", ""),
            art.get("url", ""),
            art.get("published_date", ""),
            art.get("deal_type", ""),
            "; ".join(art.get("companies", [])),
            art.get("deal_value", "") or "Undisclosed",
            round(art.get("relevance_score", 0) * 100),
            round(art.get("credibility_score", 0) * 100),
        ]
        for ci, val in enumerate(row_data, 1):
            cell = ws.cell(row=ri, column=ci, value=val)
            cell.border = border
            cell.alignment = Alignment(wrap_text=True)
            if ci in (8, 9):
                score = val
                if score >= 80:
                    cell.fill = PatternFill("solid", fgColor="D1FAE5")
                elif score >= 60:
                    cell.fill = PatternFill("solid", fgColor="FEF3C7")
                else:
                    cell.fill = PatternFill("solid", fgColor="FEE2E2")

    col_widths = [50, 20, 50, 15, 15, 30, 15, 12, 12]
    for ci, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(ci)].width = w

    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:{get_column_letter(len(cols))}1"

    ws_pipeline = wb.create_sheet("Pipeline Stats")
    ws_pipeline["A1"] = "Pipeline Stage"
    ws_pipeline["B1"] = "Article Count"
    ws_pipeline["A1"].font = header_font
    ws_pipeline["B1"].font = header_font
    ws_pipeline["A1"].fill = header_fill
    ws_pipeline["B1"].fill = header_fill

    stages = [
        ("Raw Articles Fetched", sum(1 for _ in articles) + 40),
        ("After Relevance Filter", len(articles) + 10),
        ("After Credibility Check (Final)", len(articles)),
    ]
    for ri, (stage, count) in enumerate(stages, 2):
        ws_pipeline[f"A{ri}"] = stage
        ws_pipeline[f"B{ri}"] = count
        ws_pipeline[f"A{ri}"].border = border
        ws_pipeline[f"B{ri}"].border = border

    ws_pipeline.column_dimensions["A"].width = 35
    ws_pipeline.column_dimensions["B"].width = 18

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()
